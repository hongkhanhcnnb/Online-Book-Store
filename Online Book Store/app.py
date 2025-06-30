from flask import Flask,jsonify,request,render_template,redirect,url_for,session
from flask_mysqldb import MySQL
import MySQLdb.cursors,re,datetime,os
from os import getenv
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
from decimal import Decimal
import datetime
from flask import Flask, request, send_file

from io import BytesIO
    

from utils.home import *
from utils.loginregister import *
from utils.book import *
from utils.search import *
from utils.user import *
from utils.orders import *

load_dotenv()
mysql_host = getenv('MYSQL_HOST',None)
mysql_user = getenv('MYSQL_USER',None)
mysql_password = getenv('MYSQL_PASSWORD',None)
mysql_db = getenv('MYSQL_DB',None)

app = Flask(__name__)

app.config['MYSQL_HOST'] = mysql_host
app.config['MYSQL_USER'] = mysql_user
app.config['MYSQL_PASSWORD'] = mysql_password
app.config['MYSQL_DB'] = mysql_db

app.secret_key = b'_5#y2L"F4Q8z\n\xec]/'  # all the session data is encrypted in the server so we need a secret key to encrypt and decrypt the data


mysql = MySQL(app)

def is_admin():
    # Kiểm tra xem người dùng đã đăng nhập và loại tài khoản là admin
    return session.get("accountType") == "admin"

# home page route
@app.route("/")
def homeRoute():
    booksData = allBooks(mysql)
    genreData = allGenre(mysql)
    return render_template("home.html",booksData=booksData,genreData=genreData)

# home page for customers
@app.route("/customerindex",methods=["POST","GET"])
def customerindexRoute():
    booksData = allBooks(mysql)
    genreData = allGenre(mysql)
    return render_template("customerindex.html",booksData=booksData,genreData=genreData)

# home page for admins
@app.route("/adminindex",methods=["POST","GET"])
def adminindexRoute():
    booksData = allBooks(mysql)
    genreData = allGenre(mysql)
    return render_template("adminindex.html",booksData=booksData,genreData=genreData)

# Customer Registration route
@app.route("/register",methods=["POST","GET"])
def registerRoute():
    if request.method == "POST":
        username = str(request.form.get("username"))
        fname = str(request.form.get("fname"))
        lname = str(request.form.get("lname"))
        email = str(request.form.get("email"))
        password = str(request.form.get("password"))
        phone = str(request.form.get("phone"))
        country = str(request.form.get("country"))
        state = str(request.form.get("state"))
        pincode = str(request.form.get("pincode"))
        address = str(request.form.get("address"))

        response = register(mysql,username,fname,lname,email,password,phone,country,state,pincode,address)
        
        if response == 1: # regsitration is successful
            return render_template("login.html",response=response)
        else: # registration failed
            return render_template("register.html",response=response)

    return render_template("register.html")

# login for customers and admins route
@app.route("/login",methods=["POST","GET"])
def loginRoute():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        account = request.form.get("account")

        if account=="customer":
            response = customerLogin(mysql,username,password,account)
            if response == 1: # login success
                session["userID"] = username # creating a session of the username
                session["accountType"] = account # creating a session of the account type
                return redirect(url_for("customerindexRoute"))
            else: # Login failed
                return render_template("login.html",response = response)

        if account=="admin":
            response = adminLogin(mysql,username,password,account)
            if response == 1: # login success
                session["userID"] = username # creating a session of the username
                session["accountType"] = account # creating a session of the account type
                return redirect(url_for("adminindexRoute"))
            else: # login failed
                return render_template("login.html",response=response)

    return render_template("login.html")

# search books in admin portal
@app.route("/search",methods=["POST","GET"])
def searchRoute():
    if request.method == "POST":
        search = str(request.form.get("search"))
        query = str(request.form.get("query"))

        if search == "title": # search by title
            booksData = searchTitle(mysql,query)
            return render_template("search.html",booksData=booksData,search=search)
        
        if search == "genre": # search by genre
            booksData = searchGenre(mysql,query)
            return render_template("search.html",booksData=booksData,search=search)
        
        if search == "author": # search by author
            booksData = searchAuthor(mysql,query)
            return render_template("search.html",booksData=booksData,search=search)

        return render_template("search.html")
    
    return render_template("search.html")




# search books in admin portal
@app.route('/customersearch', methods=["POST","GET"])
def customersearchRoute():
    query = request.form.get('query', '')
    search_type = request.form.get('Tìm kiếm', '')

    if search_type == "title":
        booksData = searchTitle(mysql, query)  # Tìm kiếm theo tên sách
    elif search_type == "genre":
        booksData = searchGenre(mysql, query)  # Tìm kiếm theo thể loại
    elif search_type == "author":
        booksData = searchAuthor(mysql, query)  # Tìm kiếm theo tác giả
    else:
        booksData = []

    return render_template(
        'customersearch.html',
        booksData=booksData,
        search=search_type
    )

# Add/Delete/Update Book Route for Admin
@app.route("/books",methods=["POST","GET"])
def booksRoute():
        booksData = allBooks(mysql)
        genreData = allGenre(mysql)
        return render_template("books.html",booksData=booksData,genreData=genreData)

# Add Book Route
@app.route("/addBook", methods=["POST", "GET"])
def addBookRoute():
    if request.method == "POST":
        bookID = int(request.form.get("bookID"))
        title = str(request.form.get("title"))
        genre = str(request.form.get("genre"))
        fname = str(request.form.get("fname"))
        lname = str(request.form.get("lname"))
        year = int(request.form.get("year"))
        purchase_price = float(request.form.get("purchase_price"))
        selling_price = float(request.form.get("selling_price"))
        country = str(request.form.get("country"))
        stock = int(request.form.get("stock"))

        # Database transaction
        cur = mysql.connection.cursor()
        try:
            # Add author if not exists
            cur.execute("SELECT authorID FROM Authors WHERE firstName = %s AND lastName = %s", (fname, lname))
            author = cur.fetchone()
            if not author:
                cur.execute("INSERT INTO Authors(firstName, lastName) VALUES (%s, %s)", (fname, lname))
                cur.execute("SELECT authorID FROM Authors WHERE firstName = %s AND lastName = %s", (fname, lname))
                author = cur.fetchone()
            authorID = author[0]

            # Add publisher if not exists
            cur.execute("SELECT publisherID FROM Publishers WHERE country = %s", (country,))
            publisher = cur.fetchone()
            if not publisher:
                cur.execute("INSERT INTO Publishers(country) VALUES (%s)", (country,))
                cur.execute("SELECT publisherID FROM Publishers WHERE country = %s", (country,))
                publisher = cur.fetchone()
            publisherID = publisher[0]

            # Add book to Books table
            cur.execute(
                """
                INSERT INTO Books(
                    bookID, authorID, publisherID, title, genre, publicationYear,
                    Purchase_Price, Selling_Price, Current_Stock, Quantity, Author, Category
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """,
                (bookID, authorID, publisherID, title, genre, year, purchase_price, selling_price, stock, stock, f"{fname} {lname}", genre)
            )

            # Add book to Inventory table
            cur.execute(
                "INSERT INTO Inventory(bookID, totalStock, soldStock) VALUES (%s, %s, %s)",
                (bookID, stock, 0)
            )

            mysql.connection.commit()
            response = 1  # Success
        except Exception as e:
            print("Error adding book or updating inventory:", e)
            mysql.connection.rollback()
            response = 0  # Failure
        finally:
            cur.close()

        # Return updated data to the template
        booksData = allBooks(mysql)
        genreData = allGenre(mysql)
        return render_template("books.html", booksData=booksData, genreData=genreData, response=response)

    return redirect(url_for("booksRoute"))


#get book selling price
@app.route('/get_book_price', methods=['GET'])
def get_book_price():
    book_id = request.args.get('bookID')
    cur = mysql.connection.cursor()
    cur.execute("SELECT Selling_Price FROM Books WHERE bookID = %s", (book_id,))
    result = cur.fetchone()
    cur.close()

    if result:
        return jsonify({"success": True, "price": result[0]})
    else:
        return jsonify({"success": False, "message": "Book not found"})



# Update Book Route
@app.route("/updateBook",methods=["POST","GET"])
def updateBookRoute():
    if request.method == "POST":
        bookID = str(request.form.get("bookID"))
        price2 = str(request.form.get("price2"))

        response = updateBook(mysql,bookID,price2)
        if response == 1: # book updated successfully
            booksData = allBooks(mysql)
            genreData = allGenre(mysql)
            return render_template("books.html",booksData=booksData,genreData=genreData,response=response)

        else: # book failed to update
            booksData = allBooks(mysql)
            genreData = allGenre(mysql)
            return render_template("books.html",booksData=booksData,genreData=genreData,response=response)

    return redirect(url_for("booksRoute"))

# delete book Route
@app.route("/deleteBook",methods=["POST","GET"])
def deleteBookRoute():
    if request.method == "POST":
        bookID = str(request.form.get("bookID"))
        fname = str(request.form.get("fname"))
        lname = str(request.form.get("lname"))
        country = str(request.form.get("country"))

        response = deleteBook(mysql,bookID,fname,lname,country)
        if response == 1: # book deleted successfully
            booksData = allBooks(mysql)
            genreData = allGenre(mysql)
            return render_template("books.html",booksData=booksData,genreData=genreData,response=response)

        else: # book failed to delete
            booksData = allBooks(mysql)
            genreData = allGenre(mysql)
            return render_template("books.html",booksData=booksData,genreData=genreData,response=response)

    return redirect(url_for("booksRoute"))

# display book details route for customers
@app.route("/bookdetail<subject>",methods=["POST","GET"])
def bookDetailsRoute(subject):
    bookData = bookDetail(mysql,subject)
    return render_template("bookdetail.html",bookData=bookData)

# display book details route for admin
@app.route("/bookDetailsAdmin<subject>",methods=["POST","GET"])
def bookDetailsAdminRoute(subject):
    bookData = bookDetail(mysql,subject)
    return render_template("bookdetail2.html",bookData=bookData)

# inventory route
@app.route("/inventory",methods=["POST","GET"])
def inventoryRoute():
    bookData = inventory(mysql)
    return render_template("inventory.html",bookData=bookData)

# buy book route
@app.route("/buyBook<bookID>",methods=["POST","GET"])
def buyBookRoute(bookID):
    if request.method =="POST":
        quantity = str(request.form.get("quantity"))
        bookData = totalBookPrice(mysql,bookID,quantity)
        totalPrice = int(bookData[1]) * int(quantity)
        return render_template("payment.html", bookData=bookData, quantity=quantity, totalPrice=totalPrice, float=float)
        
    return "USE POST METHOD ONLY"
    
# pay order route
@app.route("/pay<isbn>/<quantity>/<total>",methods=["POST","GET"])
def payRoute(isbn,quantity,total):
    if request.method =="POST":
        pay = str(request.form.get("pay"))

        response = orders(mysql,isbn,quantity,total,pay,session["userID"])
        return redirect(url_for('orderconfirmationRoute',response = response))
        # return render_template("orderconfirmation.html",response=response)

    return "USE POST METHOD ONLY"

# order confirmation route
@app.route("/orderconfirmation<response>",methods=["POST","GET"])
def orderconfirmationRoute(response):
    return render_template("orderconfirmation.html",response=response)


# display users route
@app.route("/users",methods=["POST","GET"])
def usersRoute():
    adminData = admin(mysql)
    customerData = customers(mysql)
    return render_template("users.html",adminData=adminData,customerData=customerData)

# display  orders in customers and admins account account
@app.route("/myorders",methods=["POST","GET"])
def ordersRoute():
    userID = session["userID"]
    accountType = session["accountType"]

    if session["accountType"] == None or session["userID"]== None:
        return "ERROR"

    if session["accountType"]=="admin":
        Data = allorders(mysql,userID)
        return render_template("myorders.html",Data=Data,accountType=accountType)

    if session["accountType"]=="customer":
        Data = myorder(mysql,userID)
        return render_template("myorders.html",Data=Data,accountType=accountType)
    
    return "ERROR"

# display logged in users account
@app.route("/myaccount",methods=["POST","GET"])
def myAccountRoute():
    userID = session["userID"]
    accountType = session["accountType"]

    if session["accountType"] == None or session["userID"]== None:
        return "ERROR"

    if session["accountType"]=="admin":
        Data = adminAccount(mysql,userID)
        return render_template("myaccount.html",Data=Data,accountType=accountType)

    if session["accountType"]=="customer":
        Data = customerAccount(mysql,userID)
        return render_template("myaccount.html",Data=Data,accountType=accountType)
    
    return "ERROR"

# contact us route
@app.route("/contactUs",methods=["POST","GET"])
def contactUsRoute():
    if request.method == "POST":
        fname = str(request.form.get("fname"))
        lname = str(request.form.get("lname"))
        email = str(request.form.get("email"))
        message = str(request.form.get("message"))
        timestamp = datetime.datetime.now()
        response = contactUs(mysql,fname,lname,email,message,timestamp)
        if response == 1:
            return "Message Submitted"
        else:
            return "Failed to add message"
            
    return "Use POST METHOD ONLY"

# logout route
@app.route("/logout",methods = ["GET","POST"])
def logoutRoute():
    session.pop("userID",None) # removing username from session variable
    session.pop("accountType",None) # removing account from session variable
    booksData = allBooks(mysql)
    genreData = allGenre(mysql)
    return render_template("home.html",booksData=booksData,genreData=genreData)

# Thêm route mới cho hủy đơn hàng
@app.route("/cancelOrder/<orderID>", methods=["POST"])
def cancelOrderRoute(orderID):
    if "userID" not in session:
        return redirect(url_for("loginRoute"))
        
    response = cancelOrder(mysql, orderID)
    return render_template("cancelconfirmation.html", response=response)

# @app.route("/inventory-report", methods=["GET", "POST"])
# def inventory_report():
#     if not is_admin():
#         return "Access Denied: Admins Only", 403

#     if request.method == "POST":
#         # B1: Nhận thông tin tháng và năm
#         month = int(request.form.get("month"))
#         year = int(request.form.get("year"))

#         # B2: Kết nối cơ sở dữ liệu
#         cur = mysql.connection.cursor()

#         # B3: Lấy dữ liệu từ bảng Inventory và Books
#         query = """
#         SELECT
#             b.bookID,
#             b.title,
#             i.totalStock AS Opening_Stock,
#             (i.totalStock - i.soldStock) AS Transactions,
#             i.soldStock AS Closing_Stock
#         FROM
#             Books b
#         INNER JOIN
#             Inventory i ON b.bookID = i.bookID;
#         """
#         cur.execute(query)
#         inventory_data = cur.fetchall()

#         # B4: Lưu báo cáo tồn kho vào InventoryReport
#         for row in inventory_data:
#             bookID, title, opening_stock, transactions, closing_stock = row
#             insert_query = """
#             INSERT INTO InventoryReport (Month, bookID, Opening_Stock, Transactions, Closing_Stock)
#             VALUES (%s, %s, %s, %s, %s)
#             ON DUPLICATE KEY UPDATE
#                 Opening_Stock = VALUES(Opening_Stock),
#                 Transactions = VALUES(Transactions),
#                 Closing_Stock = VALUES(Closing_Stock);
#             """
#             cur.execute(insert_query, (f"{year}-{month}-01", bookID, opening_stock, transactions, closing_stock))

#         # B5: Đóng kết nối
#         mysql.connection.commit()
#         cur.close()

#         # B6: Chuyển hướng tới trang hiển thị báo cáo
#         return redirect(url_for("inventory_overview", month=month, year=year))

#     # Trả về giao diện nhập thông tin báo cáo
#     return render_template("inventory-report.html")

# @app.route("/inventory-overview", methods=["GET"])
# def inventory_overview():
#     if not is_admin():
#         return "Access Denied: Admins Only", 403

#     # Nhận thông tin tháng và năm
#     month = request.args.get("month", type=int)
#     year = request.args.get("year", type=int)

#     if not month or not year:
#         return "Invalid request: Please provide both month and year.", 400

#     # Kết nối cơ sở dữ liệu
#     cur = mysql.connection.cursor()

#     # Truy vấn dữ liệu từ InventoryReport
#     query = """
#     SELECT
#         ir.Month,
#         ir.bookID,
#         b.title,
#         ir.Opening_Stock,
#         ir.Transactions,
#         ir.Closing_Stock
#     FROM
#         InventoryReport ir
#     INNER JOIN
#         Books b ON ir.bookID = b.bookID
#     WHERE
#         MONTH(ir.Month) = %s AND YEAR(ir.Month) = %s
#     """
#     cur.execute(query, (month, year))
#     inventory_data = cur.fetchall()

#     # Đóng kết nối cơ sở dữ liệu
#     cur.close()

#     # Trả về giao diện hiển thị dữ liệu
#     return render_template(
#         "inventory-overview.html",
#         inventory_data=inventory_data,
#         month=month,
#         year=year
#     )

@app.route("/inventory-report", methods=["GET", "POST"])
def inventory_report():
    if not is_admin():
        return "Access Denied: Admins Only", 403  # Chỉ cho phép admin truy cập

    # Dữ liệu tồn kho để hiển thị ngay khi nhập tháng và năm
    inventory_data = []
    month = None
    year = None

    if request.method == "POST":
        # Nhận thông tin tháng và năm
        month = int(request.form.get("month"))
        year = int(request.form.get("year"))

        # Kết nối cơ sở dữ liệu
        cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)  # Sử dụng DictCursor

        # Truy vấn báo cáo tồn kho
        query = """
        SELECT
            b.bookID,
            b.title,
            i.totalStock AS Opening_Stock,
            (i.totalStock - i.soldStock) AS Transactions,
            i.soldStock AS Closing_Stock
        FROM
            Books b
        INNER JOIN
            Inventory i ON b.bookID = i.bookID
        """
        cur.execute(query)
        inventory_data = cur.fetchall()

        # Đóng kết nối cơ sở dữ liệu
        cur.close()

    return render_template("inventory-report.html", inventory_data=inventory_data, month=month, year=year)


from flask import render_template, request, redirect, url_for
from flask_mysqldb import MySQL

@app.route("/debt-report", methods=["GET", "POST"])
def debt_report():
    if not is_admin():
        return "Access Denied: Admins Only", 403

    debt_data = []
    month = None
    year = None
    error_message = None

    if request.method == "POST":
        try:
            # Lấy giá trị tháng và năm từ form
            month = request.form.get("month")
            year = request.form.get("year")

            if not month or not year:
                error_message = "Please provide both month and year."
                raise ValueError("Month and year are required.")

            month = int(month)
            year = int(year)

            if month < 1 or month > 12:
                error_message = "Month must be between 1 and 12."
                raise ValueError("Invalid month value.")

            if year < 1900 or year > 2100:
                error_message = "Year must be a valid 4-digit year."
                raise ValueError("Invalid year value.")

            # Kết nối cơ sở dữ liệu và truy vấn
            cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
            query = """
            SELECT 
                dr.Month,
                dr.ID_Customer,
                CONCAT(c.firstName, ' ', c.lastName) AS customer_name,
                dr.Opening_Debt,
                dr.Transactions,
                dr.Closing_Debt
            FROM 
                DebtReport dr
            INNER JOIN 
                Customers c ON dr.ID_Customer = c.customerID
            WHERE 
                MONTH(dr.Month) = %s AND YEAR(dr.Month) = %s
            ORDER BY 
                dr.ID_Customer ASC
            """
            cur.execute(query, (month, year))
            debt_data = cur.fetchall()
            cur.close()
        except ValueError as ve:
            print(f"Input Error: {ve}")
        except Exception as e:
            error_message = "An error occurred while fetching the debt report."
            print(f"Database Error: {e}")

    return render_template(
        "debt-report.html",
        debt_data=debt_data,
        month=month,
        year=year,
        error_message=error_message
    )



    return render_template("debt-report.html", 
                           debt_data=debt_data, 
                           month=month, 
                           year=year, 
                           error_message=error_message)


@app.route("/payment_receipts", methods=["GET"])
def paymentReceiptsRoute():
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)  # Use DictCursor to fetch results as dictionaries
    query = "SELECT * FROM PaymentReceipt"
    cur.execute(query)
    receipts = cur.fetchall()
    cur.close()
    
    return render_template("payment_receipt/list.html", receipts=receipts)

@app.route("/payment_receipt/new", methods=["GET", "POST"])
def newPaymentReceiptRoute():
    if request.method == "POST":
        customer_name = request.form.get("customer_name")
        address = request.form.get("address")
        phone = request.form.get("phone")
        email = request.form.get("email")
        receipt_date = request.form.get("receipt_date")
        amount_collected = Decimal(request.form.get("amount_collected"))
        note = request.form.get("note")
        
        if amount_collected <= 0:
            return "Số tiền thu không được là số âm hoặc 0.", 400
        
        cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Check if customer exists and get their debt
        cur.execute("SELECT customerID, Debt FROM Customers WHERE CONCAT(firstName, ' ', lastName) = %s", (customer_name,))
        customer = cur.fetchone()
        if not customer:
            cur.close()
            return "Customer does not exist", 400
        
        customer_id = customer['customerID']
        debt = customer['Debt']
        
        # Check if amount collected exceeds debt
        if amount_collected > debt:
            cur.close()
            return "Số tiền thu không được vượt quá số tiền khách hàng đang nợ.", 400
        
        # Insert payment receipt
        cur.execute("""
            INSERT INTO PaymentReceipt (customer_name, address, phone, email, Receipt_Date, Amount_Collected, note, ID_Customer)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """, (customer_name, address, phone, email, receipt_date, amount_collected, note, customer_id))
        
        # Update customer debt
        new_debt = debt - amount_collected
        cur.execute("UPDATE Customers SET Debt = %s WHERE customerID = %s", (new_debt, customer_id))
        
        mysql.connection.commit()
        cur.close()
        
        return redirect(url_for("paymentReceiptsRoute"))
    
    current_date = datetime.datetime.now().strftime("%Y-%m-%d")
    return render_template("payment_receipt/new.html", current_date=current_date)

@app.route('/get_customer_debt/<customer_name>', methods=['GET'])
def get_customer_debt(customer_name):
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("SELECT Debt FROM Customers WHERE CONCAT(firstName, ' ', lastName) = %s", (customer_name,))
    customer = cur.fetchone()
    cur.close()
    
    if customer:
        return jsonify(customer)
    else:
        return jsonify({'error': 'Customer not found'}), 404

@app.route("/payment_receipt/<int:receipt_id>", methods=["GET"])
def paymentReceiptDetailRoute(receipt_id):
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)  # Use DictCursor to fetch results as dictionaries
    cur.execute("SELECT * FROM PaymentReceipt WHERE ID_Receipt = %s", (receipt_id,))
    receipt = cur.fetchone()
    cur.close()
    
    return render_template("payment_receipt/detail.html", receipt=receipt)

@app.route("/payment_receipt/edit/<int:receipt_id>", methods=["GET", "POST"])
def editPaymentReceiptRoute():
    if request.method == "POST":
        customer_name = request.form.get("customer_name")
        address = request.form.get("address")
        phone = request.form.get("phone")
        email = request.form.get("email")
        receipt_date = request.form.get("receipt_date")
        amount_collected = request.form.get("amount_collected")
        note = request.form.get("note")
        customer_id = request.form.get("ID_Customer")
        
        cur = mysql.connection.cursor()
        cur.execute("""
            UPDATE PaymentReceipt
            SET customer_name = %s, address = %s, phone = %s, email = %s, Receipt_Date = %s, Amount_Collected = %s, note = %s, ID_Customer = %s
            WHERE ID_Receipt = %s
        """, (customer_name, address, phone, email, receipt_date, amount_collected, note, customer_id, receipt_id))
        mysql.connection.commit()
        cur.close()
        
        return redirect(url_for("paymentReceiptsRoute"))
    
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("SELECT * FROM PaymentReceipt WHERE ID_Receipt = %s", (receipt_id,))
    receipt = cur.fetchone()
    cur.close()
    
    return render_template("payment_receipt/edit.html", receipt=receipt)

@app.route("/payment_receipt/delete/<int:receipt_id>", methods=["POST"])
def deletePaymentReceiptRoute(receipt_id):
    cur = mysql.connection.cursor()
    cur.execute("DELETE FROM PaymentReceipt WHERE ID_Receipt = %s", (receipt_id,))
    mysql.connection.commit()
    cur.close()
    
    return redirect(url_for("paymentReceiptsRoute"))



# Giả sử bạn có một danh sách sách (có thể thay bằng cơ sở dữ liệu thực tế)
# books = [
#     (1, 'Sách A', 100000, 'Thể loại A'),
#     (2, 'Sách B', 120000, 'Thể loại B'),
#     (3, 'Sách C', 90000, 'Thể loại C'),
#     (4, 'Sách D', 150000, 'Thể loại D'),
# ]

# customers = [
#     (1, 'Khách hàng A'),
#     (2, 'Khách hàng B'),
#     (3, 'Khách hàng C'),
# ]

# # Route chính để hiển thị trang lập hóa đơn
# @app.route('/create_invoice')
# def invoice():
#     today_date = datetime.date.today()
#     return render_template('create_invoice.html', books=books, customers=customers, today_date=today_date)

# # Route để lấy thông tin sách theo ID (sử dụng AJAX)
# @app.route('/get_book_info/<int:book_id>', methods=['GET'])
# def get_book_info(book_id):
#     # Tìm sách theo book_id
#     book = next((b for b in books if b[0] == book_id), None)
#     if book:
#         # Trả về dữ liệu sách dưới dạng JSON
#         return jsonify({
#             'genre': book[3],           # Thể loại
#             'Selling_Price': book[2],   # Đơn giá
#         })
#     else:
#         return jsonify({'error': 'Book not found'}), 404



# Route chính để hiển thị trang lập hóa đơn
@app.route('/create_invoice')
def invoice():
    today_date = datetime.date.today()
    
    # Lấy danh sách sách từ cơ sở dữ liệu
    cur = mysql.connection.cursor()
    cur.execute("SELECT bookID, title, Selling_Price, genre FROM Books")
    books = cur.fetchall()

    # Lấy danh sách khách hàng
    cur.execute("SELECT customerID, firstName lastName FROM Customers")
    customers = cur.fetchall()

    return render_template('create_invoice.html', books=books, customers=customers, today_date=today_date)

# Route để lấy thông tin sách theo ID (sử dụng AJAX)
@app.route('/get_book_info/<int:book_id>', methods=['GET'])
def get_book_info(book_id):
    # Lấy thông tin sách từ cơ sở dữ liệu
    cur = mysql.connection.cursor()
    cur.execute("SELECT genre, Selling_Price FROM Books WHERE bookID = %s", (book_id,))
    book = cur.fetchone()

    if book:
        # Trả về dữ liệu sách dưới dạng JSON
        return jsonify({
            'genre': book[0],           # Thể loại
            'Selling_Price': book[1],   # Đơn giá
        })
    else:
        return jsonify({'error': 'Book not found'}), 404




@app.route('/export_invoice', methods=['POST'])
def export_invoice():
    # Lấy dữ liệu từ form
    customer_name = request.form.get('customerName')
    phone_number = request.form.get('phoneNumber')  # Lấy số điện thoại
    date = request.form.get('date')
    books = request.form.getlist('bookID[]')
    categories = request.form.getlist('category[]')
    quantities = request.form.getlist('quantity[]')
    prices = request.form.getlist('price[]')
    total = request.form.get('total')
    paid = request.form.get('paid')
    remaining = request.form.get('remaining')

    # Tạo file Word
    document = Document()
    document.add_heading('Hóa Đơn Bán Sách', level=1)

    document.add_paragraph(f'Họ Tên Khách Hàng: {customer_name}')
    document.add_paragraph(f'Số Điện Thoại: {phone_number}')  # Hiển thị số điện thoại
    document.add_paragraph(f'Ngày Lập Hóa Đơn: {date}')

    # Thêm bảng chi tiết hóa đơn
    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Sách'
    hdr_cells[2].text = 'Thể loại'
    hdr_cells[3].text = 'Số lượng'
    hdr_cells[4].text = 'Đơn giá'

    for i, (book, category, quantity, price) in enumerate(zip(books, categories, quantities, prices), start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = book
        row_cells[2].text = category
        row_cells[3].text = quantity
        row_cells[4].text = price

    document.add_paragraph(f'Tổng tiền: {total}')
    document.add_paragraph(f'Số tiền trả: {paid}')
    document.add_paragraph(f'Còn lại: {remaining}')

    # Lưu vào memory buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Trả file về client
    return send_file(buffer, as_attachment=True, download_name="invoice.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/get_customer_info_by_name/<customer_name>', methods=['GET'])
def get_customer_info_by_name(customer_name):
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    # Use LIKE to handle spaces and special characters
    query = "SELECT address, phone, emailID AS email, Debt AS debt FROM Customers WHERE CONCAT(firstName, ' ', lastName) LIKE %s"
    cur.execute(query, (f"%{customer_name}%",))
    customer = cur.fetchone()
    cur.close()
    
    if customer:
        return jsonify(customer)
    else:
        print(f"Customer not found: {customer_name}")
        return jsonify({'error': 'Customer not found'}), 404


if __name__ == "__main__":
    app.run(debug=True)

